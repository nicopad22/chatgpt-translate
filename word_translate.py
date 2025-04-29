import openai
from docx import Document
import os
import time

# Function to rapidly translate all paragraphs, excluding empty lines    
def translate_paragraphs(element, translator):
    for paragraph in element.paragraphs:
            paragraph.text = translator(paragraph.text)

# Function to rapidly translate tables, making sure to manage 
def translate_table(table, translator):
    last_cell = ""
    for row in table.rows:
        for cell in row.cells:
            # make sure to skip merged cells, so AI doesnt reply to questions in english posed by itself.
            if cell.text != last_cell:
                translate_paragraphs(cell, translator)
                # update last cell's translation
                last_cell = cell.text

            for table in cell.tables:
                translate_table(table, translator)

# Function to couple the translation of paragraphs and tables all in one (useful to repeat in many sections, headers, footers, etc)
def translate_all(section, translator):
    translate_paragraphs(section, translator)
    for table in section.tables:
        translate_table(table, translator)

# External use function
def translate_word(root, file, outfile, translator):
    new_document = Document(root + outfile)

    # translate document paragraphs and tables
    translate_all(new_document, translator)

    # translate document headers and footers
    for section in new_document.sections:
        # check for specific case of different first-page footer.
        if section.different_first_page_header_footer:
            translate_all(section.first_page_header, translator)
            translate_all(section.first_page_footer, translator)
        
        # check for specific case of different even-page footer.
        if new_document.settings.odd_and_even_pages_header_footer:
            translate_all(section.even_page_header, translator)
            translate_all(section.even_page_footer, translator)

        translate_all(section.header, translator)
        translate_all(section.footer, translator)

    # SAVE DOCUMENT WITH NEW NAME
    new_document.save(root + outfile)

# main function
def main():

    root = "C:\\Users\\nicol\\Downloads\\CHATGPT-TRANSLATE\\Word\\deik7\\"

    out_language = input("output language [es/en]: ")

    # OpenAI API client
    # Store there your API key
    client = 0
    with open(".\\TOKEN.txt") as token_file:
        client = openai.OpenAI(api_key=[line for line in token_file][0]) 
        
    def translate(text):
        if text.strip() != "":
            if out_language == "es":
                return translate_to_spanish(text)
            elif out_language == "en":
                return translate_to_english(text)
        else:
            return ""
        
    # Function to translate text / AI interaction
    def translate_to_spanish(text):
        translation = text.replace(text.strip(), "x")

        result = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                    {"role": "system",
                    "content": "Eres un asistente de traducción. Traduces cualquier texto que recibas al español, sin realizar comentarios adicionales."},
                    {"role": "user", "content": f"{text.strip()}"}
                ],
            temperature=0.7,
        )

        translation = translation.replace("x", result.choices[0].message.content)

        print(text + " / " + translation)
        return translation

    # Function to translate text / AI interaction
    def translate_to_english(text):
        translation = text.replace(text.strip(), "x")

        result = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                    {"role": "system",
                    "content": "You are a translation assistant. You translate any text you recieve into english, without making any comments."},
                    {"role": "user", "content": f"{text.strip()}"}
                ],
            temperature=0.7,
        )

        translation = translation.replace("x", result.choices[0].message.content)

        print(text + " / " + translation)
        return translation

    # count the number of translated documents
    counter = 0
    
    files = os.listdir(root)
    word_files = []
    print(f" {len(files)} files found: ")
    for file in files:
        if file[-5:] == ".docx" and file[-18:] != ' [TRANSLATED].docx':
            word_files.append(file)
            print(" - " + file)

    for word_file in word_files:

        # File paths
        word_output = word_file[:-5] + ' [TRANSLATED].docx'

        if word_output in files:
            continue
        
        print(f"\n\n-------------------------\nTRANSLATING: \"{word_file}\"\n-------------------------\n\n")

        # make copy of original excel file to edit (keep formatting)
        os.popen(f"copy \"{root + word_file}\" \"{root + word_output}\"")
        time.sleep(0.4)
        # Initiate document objects
        original_document = Document(root + word_file)
        new_document = Document(root + word_output)

        # translate document paragraphs and tables
        translate_all(new_document, translate)

        # translate document headers and footers
        for section in new_document.sections:
            
            # check for specific case of different first-page footer.
            if section.different_first_page_header_footer:
                translate_all(section.first_page_header, translate)
                translate_all(section.first_page_footer, translate)
            
            # check for specific case of different even-page footer.
            if new_document.settings.odd_and_even_pages_header_footer:
                translate_all(section.even_page_header, translate)
                translate_all(section.even_page_footer, translate)

            translate_all(section.header, translate)
            translate_all(section.footer, translate)

        # SAVE DOCUMENT WITH NEW NAME
        new_document.save(root + word_output)
        counter += 1

    print(f"-------------------------------------------------------\n\nDONE!\n\nTranslated a total of {counter} word files in the \"{root}\" directory.\n")

if __name__ == "__main__":
    main()